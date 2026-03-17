import streamlit as st
import google.generativeai as genai
import PyPDF2
import os
import glob
import re
from docx import Document
from io import BytesIO

st.set_page_config(page_title="성공적인 진학을 위한 체계적인 학교생활기록부 분석", layout="wide")
st.title("🏫 성공적인 진학을 위한 체계적인 학교생활기록부 분석")
st.markdown("체계적이고 가독성 있게 학교생활기록부를 분석합니다.")

# 세션 상태 초기화 (기억 상실 방지)
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = None
if 'word_file' not in st.session_state:
    st.session_state.word_file = None

@st.cache_data(show_spinner=False)
def load_reference_pdfs(pdf_list):
    text = ""
    for pdf_file in pdf_list:
        with open(pdf_file, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    return text

# [수정됨] 지워졌던 PDF 파일 찾는 코드 복구
pdf_files = glob.glob("*.pdf")

with st.sidebar:
    st.header("🔑 기본 설정")
    
    # 비밀 금고에서 API 키 불러오기
    try:
        api_key = st.secrets["GEMINI_API_KEY"]
        st.success("✅ 안전한 금고에서 API 키를 불러왔습니다.")
    except:
        api_key = ""
        st.error("🚨 금고에 API 키가 없습니다!")
    
    st.markdown("---")
    st.subheader("📚 내장된 평가 기준 파일")
    if pdf_files:
        for f in pdf_files:
            st.write(f"- {f}")
    else:
        st.error("폴더에 기준 PDF 파일이 없습니다!")

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("📁 1. 학생 학교생활기록부 업로드")
    student_file = st.file_uploader("학생의 생기부 파일 (PDF)을 여기에 드래그 앤 드롭하세요.", type=["pdf"])

with col2:
    st.subheader("💡 2. 분석을 위한 학생의 특이사항 및 참고 내용")
    teacher_context = st.text_area(
        "예: '경영, 무역 계열로 진학을 희망하는 학생이야. 이 점 고려해서 분석해줘'", 
        height=150
    )
    submit_btn = st.button("↵ 초고속 AI 분석 시작 (클릭)", type="primary", use_container_width=True)

st.markdown("---")

def create_word_file(text):
    doc = Document()
    doc.add_heading('학교생활기록부 분석 결과', 0) 
    
    for line in text.split('\n'):
        if not line.strip():
            doc.add_paragraph()
            continue
            
        p = doc.add_paragraph()
            
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if not part: 
                continue
                
            if part.startswith('**') and part.endswith('**'):
                clean_text = part[2:-2]
                run = p.add_run(clean_text)
                run.bold = True
            else:
                run = p.add_run(part)
                
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

if submit_btn:
    if not api_key:
        st.error("API 키가 올바르게 설정되지 않았습니다.")
    elif not pdf_files:
        st.error("기준이 될 평가 PDF 파일이 없습니다.")
    elif not student_file:
        st.error("학생의 생기부 파일(PDF)을 업로드해 주세요.")
    else:
        status_box = st.empty()
        
        try:
            status_box.info("⏳ [진행상황 1/4] 내장된 가이드북(PDF)을 읽고 암기하는 중입니다...")
            reference_text = load_reference_pdfs(pdf_files)
            
            status_box.info("⏳ [진행상황 2/4] 업로드하신 학생의 생기부를 꼼꼼히 읽는 중입니다...")
            student_data_text = ""
            student_pdf_reader = PyPDF2.PdfReader(student_file)
            for page in student_pdf_reader.pages:
                text = page.extract_text()
                if text:
                    student_data_text += text + "\n"
            
            status_box.warning("🔍 [진행상황 3/4] 최적의 구글 AI 모델을 탐색 중입니다...")
            
            genai.configure(api_key=api_key)
            best_model_name = ""
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    best_model_name = m.name.replace("models/", "")
                    if 'flash' in best_model_name or 'pro' in best_model_name:
                        break 
            
            if best_model_name == "":
                raise Exception("사용 가능한 AI 모델이 없습니다.")
            
            status_box.success(f"🤖 [진행상황 4/4] '{best_model_name}' 엔진으로 맹렬히 분석 중입니다...")
            
            model = genai.GenerativeModel(best_model_name)
            
            prompt = f"""
            당신은 20년 경력의 대한민국 최고 수석 진학 상담 교사입니다.
            아래에 제공된 [대학 평가 기준 자료]를 완벽하게 숙지한 뒤, 
            이를 바탕으로 [업로드된 학생의 생기부 내용]을 날카롭고 구체적으로 분석해 주세요.

            🚨 [매우 중요한 작성 규칙 - 반드시 지킬 것!] 🚨
            1. 🌟 폰트 크기 통일을 위한 마크다운 헤더(#) 사용 절대 금지 🌟
               - 복사 및 붙여넣기 시 글자 크기가 제멋대로 커지는 것을 막기 위해, `###`나 `####` 같은 샵 기호는 절대 사용하지 마세요.
               - 제목이나 강조할 부분은 오직 양쪽 별표 두 개(`**`)를 사용해 '진하게'만 표시하세요. (예: `**1. 전공 적합성 및 주요 강점**`)
            
            2. 가독성을 위한 '3단 구조 레이아웃' (샵 기호 없이 적용):
               - (1) 테마 소제목: `**■ 데이터 분석 역량**` 형태로 작성하세요.
               - (2) 통합 출처: 다음 줄에 `**▶ [1학년 진로활동, 3학년 경제]**` 형태로 표기하세요.
               - (3) 본문 작성: 다음 줄에 3~4문장으로 압축된 본문을 작성하세요. (문장 끝 꼬리표 `[1학년 진로활동]` 필수)
               
            3. 전 구간 완벽한 개조식 어미 사용 (절대 규칙):
               - 1번부터 4번 항목까지 단 하나의 예외도 없이 모든 문장의 끝은 명사형 종결어미('~함', '~임', '~됨', '~필요함')로 끝나야 합니다. ('~다', '~합니다' 절대 금지)
            
            4. 특정 대학명 언급 완벽 차단:
               - 특정 대학교 이름은 절대 출력하지 마세요. 모두 '목표 대학'으로 통일하세요.

            💡 [완벽한 작성 예시 - 이 3단 줄바꿈 구조와 강조 기호를 반드시 똑같이 모방하세요!]
            **1. 전공 적합성 및 주요 강점**
            
            **■ 데이터 기반 비즈니스 분석 및 문제 해결 역량**
            **▶ [1학년 진로활동, 3학년 자율활동, 3학년 경제수학]**
            학생은 컴퓨터 정규 교과 및 동아리 활동을 통해 비즈니스 의사결정 역량을 꾸준히 발전시켜 온 모습이 돋보임. 1학년 때 생성형 언어 모델 API 응용 프로그래밍에 참여하여 파이썬 코드를 직접 설계하며 실시간 데이터 가공 방법을 학습함 [1학년 진로활동]. 3학년 행사에서는 라즈베리파이와 센서를 활용한 달걀 낙하 실험에서 실시간 데이터를 수집하고 물리 공식을 적용해 충격량을 계산하는 등 과학적 분석 능력을 향상함 [3학년 자율활동]. 수요 함수와 공급 함수의 그래프를 해석하여 균형 가격과 거래량을 능숙하게 계산하는 등 경영학 전공에 필요한 정량적 분석 능력을 효과적으로 발전시킨 점이 돋보임 [3학년 경제수학].

            [담당 교사의 특별 지시사항 및 학생 특이사항]
            {teacher_context if teacher_context else "특별한 지시사항 없음. 일반적인 기준에 따라 철저히 분석할 것."}

            [대학 평가 기준 자료 (통합 내용)]
            {reference_text}

            [업로드된 학생의 생기부 내용]
            {student_data_text}

            반드시 위의 예시 형태 그대로, 샵(#) 기호 없이 아래 4가지 양식으로 답변해 주세요. 모든 문장은 무조건 명사형(~함, ~임)으로 끝내세요.
            **1. 전공 적합성 및 주요 강점** (**■ 소제목**, **▶ 통합출처**, 본문 구조 엄수)
            **2. 평가 기준에 비추어 볼 때 보완이 필요한 약점** (**■ 소제목**, **▶ 통합출처**, 본문 구조 엄수)
            **3. 추천 심화 탐구 주제 및 면접 예상 압박 질문 3가지**
            **4. 종합 의견 및 향후 발전 방향 (구체적인 액션 플랜 포함)**
            """
            
            response = model.generate_content(prompt)
            
            st.session_state.analysis_result = response.text
            st.session_state.word_file = create_word_file(response.text)
            
            status_box.success("✅ [분석 완료!] 초고속 심층 분석이 완료되었습니다. 결과물을 확인해 주세요!")
            
        except Exception as e:
            status_box.error(f"오류가 발생했습니다: ({e})")

if st.session_state.analysis_result:
    st.write(st.session_state.analysis_result)
    
    st.download_button(
        label="📥 학교생활기록부 분석 결과 워드(Word) 파일 다운로드 (클릭)",
        data=st.session_state.word_file,
        file_name="학교생활기록부 분석 결과.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
